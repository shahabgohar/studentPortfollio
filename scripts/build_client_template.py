from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/client-quotation-template.docx"

PRIMARY = "111827"
TEXT = "374151"
MUTED = "6B7280"
ACCENT = "565BCF"
LIGHT_ACCENT = "EEF0FF"
SOFT = "F8FAFC"
BORDER = "D9DCEB"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=140, start=140, bottom=140, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = "w:{}".format(edge)
                element = borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    borders.append(element)
                element.set(qn("w:val"), "nil")


def set_run(run, size=None, bold=False, color=None, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_label_value(paragraph, label, value):
    r = paragraph.add_run(label)
    set_run(r, 9, True, MUTED)
    r = paragraph.add_run(value)
    set_run(r, 9, False, PRIMARY)


def add_section_heading(doc, title, subtitle=None, page_break_before=False):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    run = p.add_run(title)
    set_run(run, 17, True, PRIMARY)
    if subtitle:
        s = doc.add_paragraph()
        s.paragraph_format.space_after = Pt(8)
        s.paragraph_format.keep_with_next = True
        run = s.add_run(subtitle)
        set_run(run, 9.5, False, MUTED)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run, 10, False, TEXT)


def add_placeholder_line(doc, label, value="[Add details here]"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_label_value(p, f"{label}: ", value)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.color.rgb = RGBColor.from_string(PRIMARY)

    doc.styles["Title"].font.name = "Arial"
    doc.styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("shahabgohar.dev  |  AI + Full-Stack Product Engineering")
    set_run(run, 8.5, False, MUTED)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Shahab Gohar  |  shahabgohar.dev  |  shahab.developer.work@gmail.com  |  linkedin.com/in/shahabgohar")
    set_run(run, 8, False, MUTED)

    # Cover/title block.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("SHAHAB GOHAR")
    set_run(run, 11, True, ACCENT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Client Communication & Quotation")
    set_run(run, 26, True, PRIMARY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("AI product engineering, full-stack web applications, SuiteCRM customization, and business automation.")
    set_run(run, 11, False, TEXT)

    meta = doc.add_table(rows=2, cols=4)
    set_table_width(meta, 9360)
    remove_table_borders(meta)
    widths = [1800, 2700, 1800, 3060]
    labels = [
        ("Prepared for", "[Client / Company]"),
        ("Prepared by", "Shahab Gohar"),
        ("Date", "[Month DD, YYYY]"),
        ("Quote #", "[SG-0001]"),
        ("Website", "https://shahabgohar.dev"),
        ("Email", "shahab.developer.work@\ngmail.com"),
        ("Validity", "[14 days]"),
        ("Project", "[Project name]"),
    ]
    for i, cell in enumerate(meta._cells):
        set_cell_width(cell, widths[i % 4])
        set_cell_margins(cell, 130, 130, 130, 130)
        set_cell_shading(cell, LIGHT_ACCENT if i < 4 else SOFT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(labels[i][0] + "\n")
        set_run(r, 7.5, True, MUTED)
        r = p.add_run(labels[i][1])
        set_run(r, 9.2, False, PRIMARY)

    add_section_heading(doc, "Executive Summary", "Use this section to make the client feel understood before showing price.")
    p = doc.add_paragraph()
    run = p.add_run(
        "Thank you for considering me for this project. Based on our discussion, the objective is to "
        "[describe the business goal], reduce friction in [workflow], and deliver a polished product experience "
        "that is reliable enough for real users."
    )
    set_run(run, 10.5, False, TEXT)

    callout = doc.add_table(rows=1, cols=1)
    set_table_width(callout, 9360)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_ACCENT)
    set_cell_border(cell, ACCENT, "10")
    set_cell_margins(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Recommended direction: ")
    set_run(r, 10, True, ACCENT)
    r = p.add_run("[Short recommendation: MVP, automation workflow, AI assistant, SuiteCRM customization, or full-stack build.]")
    set_run(r, 10, False, PRIMARY)

    add_section_heading(doc, "Scope of Work")
    for item in [
        "Discovery and requirement clarification for the target workflow.",
        "UI/UX implementation aligned with the client brand and user journey.",
        "Backend/API integration, database work, automation logic, or SuiteCRM customization as required.",
        "Testing, bug fixing, responsive checks, and deployment handoff.",
        "Documentation or walkthrough notes for using and maintaining the delivered work.",
    ]:
        add_bullet(doc, item)

    add_section_heading(doc, "Deliverables")
    deliverables = doc.add_table(rows=1, cols=4)
    set_table_width(deliverables, 9360)
    headers = ["#", "Deliverable", "Description", "Acceptance Check"]
    col_widths = [520, 2500, 3460, 2880]
    for idx, header_text in enumerate(headers):
        cell = deliverables.rows[0].cells[idx]
        set_cell_width(cell, col_widths[idx])
        set_cell_shading(cell, ACCENT)
        set_cell_border(cell, ACCENT)
        set_cell_margins(cell, 150, 140, 150, 140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header_text)
        set_run(r, 9, True, WHITE)
    for row_data in [
        ("1", "[Feature / Module]", "[What will be built]", "[How client confirms it is complete]"),
        ("2", "[Integration]", "[API, CRM, payment, AI, or automation detail]", "[Successful test / expected result]"),
        ("3", "[Launch / Handoff]", "[Launch, docs, training, or support]", "[Client can use it independently]"),
    ]:
        cells = deliverables.add_row().cells
        for idx, value in enumerate(row_data):
            cell = cells[idx]
            set_cell_width(cell, col_widths[idx])
            set_cell_border(cell)
            set_cell_margins(cell, 150, 140, 150, 140)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            set_run(r, 9.2, False, TEXT)

    add_section_heading(doc, "Investment")
    pricing = doc.add_table(rows=1, cols=5)
    set_table_width(pricing, 9360)
    headers = ["Item", "Description", "Qty", "Rate", "Total"]
    widths = [2200, 3360, 800, 1400, 1600]
    for idx, header_text in enumerate(headers):
        cell = pricing.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, PRIMARY)
        set_cell_border(cell, PRIMARY)
        set_cell_margins(cell, 150, 140, 150, 140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header_text)
        set_run(r, 9, True, WHITE)
    for row_data in [
        ("Discovery & planning", "[Requirements, architecture, project plan]", "1", "[PKR/USD]", "[Amount]"),
        ("Development", "[Frontend, backend, CRM, AI, or automation work]", "[Qty]", "[Rate]", "[Amount]"),
        ("Testing & handoff", "[QA, deployment support, documentation]", "1", "[Rate]", "[Amount]"),
        ("Total", "", "", "", "[Total Amount]"),
    ]:
        cells = pricing.add_row().cells
        for idx, value in enumerate(row_data):
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_border(cell)
            set_cell_margins(cell, 150, 140, 150, 140)
            if row_data[0] == "Total":
                set_cell_shading(cell, LIGHT_ACCENT)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run(r, 9.2, row_data[0] == "Total", PRIMARY if row_data[0] == "Total" else TEXT)

    break_paragraph = doc.add_paragraph()
    break_paragraph.paragraph_format.space_before = Pt(0)
    break_paragraph.paragraph_format.space_after = Pt(0)
    break_run = break_paragraph.add_run()
    set_run(break_run, 1, False, WHITE)
    break_run.add_break(WD_BREAK.PAGE)
    add_section_heading(doc, "Timeline & Milestones")
    timeline = doc.add_table(rows=1, cols=3)
    set_table_width(timeline, 9360)
    headers = ["Milestone", "Target Date", "Notes"]
    widths = [2900, 1800, 4660]
    for idx, header_text in enumerate(headers):
        cell = timeline.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, ACCENT)
        set_cell_border(cell, ACCENT)
        set_cell_margins(cell, 150, 140, 150, 140)
        r = cell.paragraphs[0].add_run(header_text)
        set_run(r, 9, True, WHITE)
    for row_data in [
        ("Kickoff / access", "[Date]", "[Client provides access, references, credentials, assets]"),
        ("First working version", "[Date]", "[Review core workflow and priority screens]"),
        ("Final delivery", "[Date]", "[QA, revisions, deployment, handoff]"),
    ]:
        cells = timeline.add_row().cells
        for idx, value in enumerate(row_data):
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_border(cell)
            set_cell_margins(cell, 150, 140, 150, 140)
            r = cell.paragraphs[0].add_run(value)
            set_run(r, 9.2, False, TEXT)

    add_section_heading(doc, "Client Responsibilities")
    for item in [
        "Provide required access, brand assets, existing documentation, and example workflows.",
        "Review milestone deliveries within [X business days] to keep the schedule on track.",
        "Confirm final acceptance after the agreed deliverables pass the acceptance checks.",
    ]:
        add_bullet(doc, item)

    add_section_heading(doc, "Terms & Notes")
    for item in [
        "Quote validity: [14 days unless otherwise stated].",
        "Payment terms: [e.g., 50% upfront, 50% before final handoff].",
        "Out-of-scope changes will be estimated separately before implementation.",
        "Post-delivery support: [X days] for bug fixes related to the agreed scope.",
    ]:
        add_bullet(doc, item)

    add_section_heading(doc, "Next Step")
    p = doc.add_paragraph()
    r = p.add_run(
        "If this proposal looks good, reply with approval and I will prepare the kickoff checklist, project workspace, "
        "and first milestone plan."
    )
    set_run(r, 10.5, False, TEXT)

    sig = doc.add_table(rows=1, cols=2)
    set_table_width(sig, 9360)
    remove_table_borders(sig)
    for idx, cell in enumerate(sig.rows[0].cells):
        set_cell_width(cell, 4680)
        set_cell_margins(cell, 180, 0, 100, 0)
        p = cell.paragraphs[0]
        label = "Client approval" if idx == 0 else "Shahab Gohar"
        r = p.add_run(label + "\n\n")
        set_run(r, 9, True, PRIMARY)
        r = p.add_run("Signature: __________________________\nDate: _______________________________")
        set_run(r, 9, False, MUTED)

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
